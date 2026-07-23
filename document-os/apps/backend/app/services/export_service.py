"""Export business logic: render documents to markdown/html/pdf/docx/json files."""
import json
import re
from datetime import datetime, timezone
from html import escape as html_escape
from pathlib import Path

import markdown as md_lib
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.errors import AppError
from app.models import Document, DocumentSection, Export
from app.repositories import document_repo, export_repo
from app.services import document_service
from app.utils.markdown import sections_to_markdown, slugify

_EXTENSIONS = {"markdown": "md", "html": "html", "pdf": "pdf", "docx": "docx", "json": "json"}


def export_document(
    db: Session, document: Document, format: str, summary: str | None = None
) -> Export:
    """Render a document to a real file and record the export."""
    fmt = format.lower()
    if fmt not in _EXTENSIONS:
        raise AppError(f"Unsupported export format: {format}")
    export_dir = Path(get_settings().EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)

    export = export_repo.create(
        db,
        obj_in={"document_id": document.id, "format": fmt, "file_path": "", "status": "completed"},
    )
    filename = f"{slugify(document.title)}-{export.id}.{_EXTENSIONS[fmt]}"
    file_path = export_dir / filename

    sections = document_service.get_sections(db, document.id)
    if fmt == "markdown":
        _write_markdown(file_path, document, sections, summary)
    elif fmt == "html":
        _write_html(file_path, document, sections, summary)
    elif fmt == "pdf":
        _write_pdf(file_path, document, sections, summary)
    elif fmt == "docx":
        _write_docx(file_path, document, sections, summary)
    else:
        _write_json(file_path, document, sections)

    export = export_repo.update(db, db_obj=export, obj_in={"file_path": str(file_path.resolve())})
    document_repo.update(db, db_obj=document, obj_in={"status": "exported"})
    return export


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def _write_markdown(
    file_path: Path, document: Document, sections: list[DocumentSection], summary: str | None
) -> None:
    """Write the full markdown export, optionally prefixed with a summary block."""
    body = sections_to_markdown(sections, document.title)
    if summary:
        quoted = "\n".join(f"> {line}" for line in summary.strip().splitlines())
        body = f"> Executive summary\n>\n{quoted}\n\n{body}"
    file_path.write_text(body, encoding="utf-8")


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

_HTML_SHELL = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
         max-width: 760px; margin: 48px auto; padding: 0 24px; line-height: 1.65; color: #1f2430; }}
  header.doc-header {{ border-bottom: 2px solid #e5e7eb; margin-bottom: 32px; padding-bottom: 16px; }}
  header.doc-header h1 {{ margin: 0 0 4px; font-size: 2rem; }}
  header.doc-header .date {{ color: #6b7280; font-size: 0.9rem; }}
  .summary-callout {{ background: #eef2ff; border-left: 4px solid #6366f1; border-radius: 4px;
                     padding: 12px 16px; margin: 0 0 32px; }}
  .summary-callout strong {{ display: block; margin-bottom: 4px; color: #4338ca; }}
  h1, h2, h3 {{ line-height: 1.3; margin: 1.6em 0 0.6em; }}
  h1 {{ font-size: 1.7rem; border-bottom: 1px solid #e5e7eb; padding-bottom: 6px; }}
  h2 {{ font-size: 1.4rem; border-bottom: 1px solid #f0f1f3; padding-bottom: 4px; }}
  h3 {{ font-size: 1.15rem; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; font-size: 0.95rem; }}
  th, td {{ border: 1px solid #d1d5db; padding: 8px 12px; text-align: left; vertical-align: top; }}
  th {{ background: #f3f4f6; font-weight: 600; }}
  tr:nth-child(even) td {{ background: #fafafa; }}
  code {{ background: #f3f4f6; border-radius: 4px; padding: 2px 5px;
         font-family: "SF Mono", Consolas, "Liberation Mono", Menlo, monospace; font-size: 0.88em; }}
  pre {{ background: #f6f7f9; border: 1px solid #e5e7eb; border-radius: 6px;
        padding: 14px 16px; overflow-x: auto; }}
  pre code {{ background: none; padding: 0; }}
  blockquote {{ border-left: 4px solid #d1d5db; margin: 1em 0; padding: 4px 16px; color: #4b5563; }}
  a {{ color: #4f46e5; }}
  @media print {{ body {{ margin: 0; max-width: none; }} }}
</style>
</head>
<body>
<header class="doc-header">
  <h1>{title}</h1>
  <div class="date">Exported {date}</div>
</header>
{summary_block}
<main>
{content}
</main>
</body>
</html>
"""


def _write_html(
    file_path: Path, document: Document, sections: list[DocumentSection], summary: str | None
) -> None:
    """Write a self-contained, print-friendly HTML export."""
    body_md = sections_to_markdown(sections)
    content = md_lib.markdown(body_md, extensions=["tables", "fenced_code", "toc"])
    summary_block = ""
    if summary:
        summary_block = (
            '<div class="summary-callout"><strong>Executive summary</strong>'
            f"{html_escape(summary.strip())}</div>"
        )
    html = _HTML_SHELL.format(
        title=html_escape(document.title),
        date=datetime.now(timezone.utc).strftime("%B %d, %Y"),
        summary_block=summary_block,
        content=content,
    )
    file_path.write_text(html, encoding="utf-8")


# ---------------------------------------------------------------------------
# Shared markdown line tokenizer (used by PDF and DOCX renderers)
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")
_TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-{2,}:?$")


def _tokenize_markdown(markdown_text: str) -> list[tuple]:
    """Tokenize markdown into (kind, payload) blocks: heading/bullet/numbered/code/table/paragraph."""
    blocks: list[tuple] = []
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            blocks.append(("code", code_lines))
        elif stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [_split_table_row(l) for l in table_lines]
            rows = [r for r in rows if not all(_TABLE_SEPARATOR_CELL_RE.match(c) for c in r)]
            if rows:
                blocks.append(("table", rows))
        elif m := _HEADING_RE.match(line):
            blocks.append(("heading", (len(m.group(1)), m.group(2).strip())))
        elif m := _BULLET_RE.match(line):
            blocks.append(("bullet", m.group(1).strip()))
        elif m := _NUMBERED_RE.match(line):
            blocks.append(("numbered", (m.group(1), m.group(2).strip())))
        elif stripped:
            blocks.append(("paragraph", stripped))
        i += 1
    return blocks


def _split_table_row(line: str) -> list[str]:
    """Split a pipe-table row into trimmed cell strings."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _inline_markup(text: str) -> str:
    """Escape text and convert **bold**/*italic* to reportlab mini-HTML."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return re.sub(r"\*([^*\n]+?)\*", r"<i>\1</i>", text)


# ---------------------------------------------------------------------------
# PDF (reportlab platypus)
# ---------------------------------------------------------------------------

def _pdf_styles() -> dict:
    """Build the paragraph style set used by the PDF renderer."""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle

    styles = {
        "title": ParagraphStyle("DocTitle", fontName="Helvetica-Bold", fontSize=24, spaceAfter=8),
        "date": ParagraphStyle("DocDate", fontName="Helvetica", fontSize=10,
                               textColor=colors.HexColor("#6b7280"), spaceAfter=4),
        "summary": ParagraphStyle("DocSummary", fontName="Helvetica-Oblique", fontSize=11,
                                  textColor=colors.HexColor("#374151"), spaceAfter=6,
                                  backColor=colors.HexColor("#eef2ff"), borderPadding=8),
        "body": ParagraphStyle("Body", fontName="Helvetica", fontSize=10.5,
                               leading=15, spaceAfter=8),
        "bullet": ParagraphStyle("Bullet", fontName="Helvetica", fontSize=10.5,
                                 leading=15, leftIndent=18, spaceAfter=3),
        "code": ParagraphStyle("Code", fontName="Courier", fontSize=8.5, leading=11,
                               backColor=colors.HexColor("#f3f4f6"), borderPadding=6,
                               spaceBefore=6, spaceAfter=8),
        "cell": ParagraphStyle("Cell", fontName="Helvetica", fontSize=9, leading=12),
    }
    for level, size in ((1, 20), (2, 16), (3, 14), (4, 12), (5, 11), (6, 11)):
        styles[f"h{level}"] = ParagraphStyle(
            f"H{level}", fontName="Helvetica-Bold", fontSize=size,
            spaceBefore=14, spaceAfter=6,
        )
    return styles


def _write_pdf(
    file_path: Path, document: Document, sections: list[DocumentSection], summary: str | None
) -> None:
    """Write a PDF export rendered from the section markdown."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
        XPreformatted,
    )

    styles = _pdf_styles()
    story: list = [Paragraph(_inline_markup(document.title), styles["title"])]
    if summary:
        story.append(Paragraph(f"<b>Executive summary.</b> {_inline_markup(summary.strip())}",
                               styles["summary"]))
    story.append(Paragraph(datetime.now(timezone.utc).strftime("%B %d, %Y"), styles["date"]))
    story.append(Spacer(1, 12))

    for kind, payload in _tokenize_markdown(sections_to_markdown(sections)):
        if kind == "heading":
            level, text = payload
            story.append(Paragraph(_inline_markup(text), styles[f"h{min(level, 6)}"]))
        elif kind == "bullet":
            story.append(Paragraph(f"&bull; {_inline_markup(payload)}", styles["bullet"]))
        elif kind == "numbered":
            number, text = payload
            story.append(Paragraph(f"{number}. {_inline_markup(text)}", styles["bullet"]))
        elif kind == "code":
            story.append(XPreformatted("\n".join(payload) or " ", styles["code"]))
        elif kind == "table":
            data = []
            for row_index, row in enumerate(payload):
                cells = []
                for cell in row:
                    text = _inline_markup(cell)
                    if row_index == 0:
                        text = f"<b>{text}</b>"
                    cells.append(Paragraph(text, styles["cell"]))
                data.append(cells)
            table = Table(data, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f3f4f6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 4))
            story.append(table)
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(_inline_markup(payload), styles["body"]))

    doc = SimpleDocTemplate(
        str(file_path), pagesize=A4, title=document.title,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
    )
    doc.build(story)


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------

def _docx_add_runs(paragraph, text: str) -> None:
    """Add runs to a python-docx paragraph honoring **bold** and *italic*."""
    for bold_part in re.split(r"(\*\*.+?\*\*)", text):
        if bold_part.startswith("**") and bold_part.endswith("**") and len(bold_part) > 4:
            paragraph.add_run(bold_part[2:-2]).bold = True
        else:
            for part in re.split(r"(\*[^*\n]+?\*)", bold_part):
                if part.startswith("*") and part.endswith("*") and len(part) > 2:
                    paragraph.add_run(part[1:-1]).italic = True
                elif part:
                    paragraph.add_run(part)


def _docx_shade(paragraph, fill: str = "F3F4F6") -> None:
    """Apply a light background shading to a python-docx paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shading)


def _write_docx(
    file_path: Path, document: Document, sections: list[DocumentSection], summary: str | None
) -> None:
    """Write a DOCX export rendered from the section markdown."""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    doc.add_heading(document.title, level=0)
    if summary:
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.add_run("Executive summary. ").bold = True
        summary_paragraph.add_run(summary.strip()).italic = True

    for kind, payload in _tokenize_markdown(sections_to_markdown(sections)):
        if kind == "heading":
            level, text = payload
            doc.add_heading(text, level=min(level, 9))
        elif kind == "bullet":
            _docx_add_runs(doc.add_paragraph(style="List Bullet"), payload)
        elif kind == "numbered":
            _, text = payload
            _docx_add_runs(doc.add_paragraph(style="List Number"), text)
        elif kind == "code":
            for code_line in payload or [""]:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(code_line if code_line else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                _docx_shade(paragraph)
        elif kind == "table":
            table = doc.add_table(rows=len(payload), cols=len(payload[0]))
            table.style = "Table Grid"
            for row_index, row in enumerate(payload):
                for col_index, cell in enumerate(row):
                    cell_paragraph = table.cell(row_index, col_index).paragraphs[0]
                    run = cell_paragraph.add_run(cell)
                    if row_index == 0:
                        run.bold = True
        else:
            _docx_add_runs(doc.add_paragraph(), payload)

    doc.save(str(file_path))


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------

def _section_tree(sections: list[DocumentSection]) -> list[dict]:
    """Build a nested tree of sections (with content) from a flat ordered list."""
    by_parent: dict[str | None, list[DocumentSection]] = {}
    for section in sections:
        by_parent.setdefault(section.parent_id, []).append(section)

    def build(parent_id: str | None) -> list[dict]:
        return [
            {
                "id": section.id,
                "title": section.title,
                "status": section.status,
                "word_count": section.word_count,
                "content": section.content,
                "children": build(section.id),
            }
            for section in by_parent.get(parent_id, [])
        ]

    return build(None)


def _write_json(file_path: Path, document: Document, sections: list[DocumentSection]) -> None:
    """Write a structured JSON export with the nested section tree."""
    payload = {
        "title": document.title,
        "description": document.description,
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "tree": _section_tree(sections),
    }
    file_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
