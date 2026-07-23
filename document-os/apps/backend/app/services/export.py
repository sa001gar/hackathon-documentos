"""Export service: Markdown / JSON / HTML / DOCX / PDF.

The ExporterAgent composes canonical Markdown; this service converts it into
the requested artifact and stores it under EXPORT_DIR.
"""
import html as html_lib
import json
import re
import uuid
from pathlib import Path

import markdown as md_lib
from sqlalchemy.orm import Session

from app.ai.agents import ExporterAgent
from app.ai.engine import get_engine
from app.core.config import get_settings
from app.core.errors import NotFoundError
from app.models import Export, User
from app.repositories import ExportRepository
from app.schemas.misc import ExportOut
from app.services.documents import DocumentService

HTML_SHELL = """<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>
  body {{ font-family: Georgia, 'Times New Roman', serif; max-width: 760px; margin: 48px auto;
         padding: 0 24px; line-height: 1.65; color: #1a1a1a; }}
  h1,h2,h3,h4 {{ font-family: -apple-system, 'Segoe UI', sans-serif; line-height: 1.25; }}
  h1 {{ border-bottom: 2px solid #eee; padding-bottom: 12px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
  th {{ background: #f6f6f6; }}
  code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 4px; font-size: 0.9em; }}
  pre code {{ display: block; padding: 16px; overflow-x: auto; }}
  blockquote {{ border-left: 3px solid #ddd; margin: 0; padding-left: 16px; color: #555; }}
</style></head><body>{body}</body></html>"""


class ExportService:
    def __init__(self, db: Session) -> None:
        self.db = db
        self.documents = DocumentService(db)
        self.exports = ExportRepository(db)

    def export_document(
        self, user: User, document_id: str, fmt: str, include_toc: bool = True
    ) -> ExportOut:
        doc = self.documents.get_owned(user, document_id)
        roots = [s for s in doc.sections if s.parent_id is None]
        agent = ExporterAgent(get_engine())
        markdown_text, _meta = agent.compose(doc, roots, include_toc=include_toc)

        export_id = uuid.uuid4().hex
        export_dir = Path(get_settings().EXPORT_DIR)
        export_dir.mkdir(parents=True, exist_ok=True)
        path = export_dir / f"{export_id}.{fmt}"

        if fmt == "md":
            path.write_text(markdown_text, encoding="utf-8")
        elif fmt == "json":
            path.write_text(self._to_json(doc, roots), encoding="utf-8")
        elif fmt == "html":
            path.write_text(self._to_html(doc.title, markdown_text), encoding="utf-8")
        elif fmt == "docx":
            self._to_docx(doc.title, markdown_text, path)
        elif fmt == "pdf":
            self._to_pdf(doc.title, markdown_text, path)
        else:  # pragma: no cover — schema validation prevents this
            raise NotFoundError(f"Unsupported format: {fmt}")

        record = Export(document_id=doc.id, format=fmt, file_path=str(path), created_by=user.id)
        record.id = export_id
        self.exports.add(record)
        self.db.commit()
        out = ExportOut.model_validate(record)
        out.download_url = f"/api/v1/exports/{record.id}/download"
        return out

    def list_for_document(self, user: User, document_id: str) -> list[ExportOut]:
        self.documents.get_owned(user, document_id)
        out = []
        for record in self.exports.for_document(document_id):
            dto = ExportOut.model_validate(record)
            dto.download_url = f"/api/v1/exports/{record.id}/download"
            out.append(dto)
        return out

    def get_file(self, user: User, export_id: str) -> tuple[Path, str]:
        record = self.exports.get(export_id)
        if not record or record.document.project.workspace.owner_id != user.id:
            raise NotFoundError("Export not found")
        path = Path(record.file_path)
        if not path.exists():
            raise NotFoundError("Export file is missing on disk")
        media = {
            "md": "text/markdown", "json": "application/json", "html": "text/html",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf",
        }
        return path, media.get(record.format, "application/octet-stream")

    # ---------------- converters ----------------

    @staticmethod
    def _to_json(doc, roots) -> str:
        def node(section):
            return {
                "title": section.title,
                "content": section.content,
                "status": section.status,
                "children": [node(c) for c in sorted(section.children, key=lambda s: s.order_index)],
            }
        payload = {
            "title": doc.title,
            "doc_type": doc.doc_type,
            "exported_from": "DocumentOS",
            "sections": [node(s) for s in sorted(roots, key=lambda s: s.order_index)],
        }
        return json.dumps(payload, indent=2)

    @staticmethod
    def _to_html(title: str, markdown_text: str) -> str:
        body = md_lib.markdown(
            markdown_text, extensions=["tables", "fenced_code", "sane_lists"]
        )
        return HTML_SHELL.format(title=html_lib.escape(title), body=body)

    @staticmethod
    def _to_docx(title: str, markdown_text: str, path: Path) -> None:
        from docx import Document as DocxDocument
        from docx.shared import Pt

        document = DocxDocument()
        document.core_properties.title = title
        in_code = False
        table_buffer: list[list[str]] = []

        def flush_table():
            if not table_buffer:
                return
            rows = [r for r in table_buffer if not all(set(c) <= set("-: ") for c in r)]
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        table.rows[i].cells[j].text = cell
            table_buffer.clear()

        for raw in markdown_text.splitlines():
            line = raw.rstrip()
            if line.strip().startswith("```"):
                in_code = not in_code
                continue
            if in_code:
                p = document.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                continue
            if line.strip().startswith("|"):
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                table_buffer.append(cells)
                continue
            flush_table()
            stripped = line.strip()
            if not stripped:
                continue
            heading = re.match(r"^(#{1,6})\s+(.*)", stripped)
            if heading:
                document.add_heading(_strip_inline(heading.group(2)), level=len(heading.group(1)))
            elif re.match(r"^[-*]\s+\[[ xX]\]\s+", stripped):
                checked = "[x]" in stripped.lower()
                text = re.sub(r"^[-*]\s+\[[ xX]\]\s+", "", stripped)
                document.add_paragraph(("☑ " if checked else "☐ ") + _strip_inline(text), style="List Bullet")
            elif re.match(r"^[-*]\s+", stripped):
                document.add_paragraph(_strip_inline(re.sub(r"^[-*]\s+", "", stripped)), style="List Bullet")
            elif re.match(r"^\d+\.\s+", stripped):
                document.add_paragraph(_strip_inline(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
            else:
                document.add_paragraph(_strip_inline(stripped))
        flush_table()
        document.save(str(path))

    @staticmethod
    def _to_pdf(title: str, markdown_text: str, path: Path) -> None:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.lib import colors

        styles = getSampleStyleSheet()
        story: list = []
        table_buffer: list[list[str]] = []
        in_code = False

        def flush_table():
            if not table_buffer:
                return
            rows = [r for r in table_buffer if not all(set(c) <= set("-: ") for c in r)]
            if rows:
                tbl = Table(rows, hAlign="LEFT")
                tbl.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.extend([tbl, Spacer(1, 0.3 * cm)])
            table_buffer.clear()

        for raw in markdown_text.splitlines():
            line = raw.rstrip()
            if line.strip().startswith("```"):
                in_code = not in_code
                continue
            if in_code:
                story.append(Paragraph(html_lib.escape(line) or "&nbsp;", styles["Code"]))
                continue
            if line.strip().startswith("|"):
                table_buffer.append([c.strip() for c in line.strip().strip("|").split("|")])
                continue
            flush_table()
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 0.15 * cm))
                continue
            heading = re.match(r"^(#{1,6})\s+(.*)", stripped)
            if heading:
                level = min(len(heading.group(1)), 3)
                story.append(Paragraph(html_lib.escape(_strip_inline(heading.group(2))), styles[f"Heading{level}"]))
            elif re.match(r"^([-*]|\d+\.)\s+", stripped):
                text = re.sub(r"^([-*]|\d+\.)\s+", "", stripped)
                story.append(Paragraph("• " + _md_inline_to_html(text), styles["BodyText"]))
            else:
                story.append(Paragraph(_md_inline_to_html(stripped), styles["BodyText"]))
        flush_table()

        pdf = SimpleDocTemplate(str(path), pagesize=A4, title=title,
                                leftMargin=2 * cm, rightMargin=2 * cm)
        pdf.build(story)


def _strip_inline(text: str) -> str:
    return re.sub(r"[*`]", "", text).strip()


def _md_inline_to_html(text: str) -> str:
    escaped = html_lib.escape(text)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", escaped)
    escaped = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", escaped)
    return escaped
